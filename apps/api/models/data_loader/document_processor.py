import io
import os
import re
import subprocess
import sys
import tempfile
from abc import ABC, abstractmethod
from typing import List

from docx import Document as WordDocument
from langchain.schema import Document
from langchain.text_splitter import CharacterTextSplitter
from loguru import logger
from models.base.dataset import Dataset, Document as DocumentModel
from models.data_loader.document_settings import (
    PDFRetrivalOption,
    PDFSplitterOption,
)
import pypdf
from utils.StorageClient import AnnotatedDataStorageClient, GoogleCloudStorageClient

# Mixins


class DocumentProcessingMixin:
    def get_text_splitter(self, document: Document) -> CharacterTextSplitter:
        options = PDFRetrivalOption(
            splitter=PDFSplitterOption(
                chunk_overlap=document.split_option.get("chunk_overlap", 0),
                chunk_size=document.split_option.get("chunk_size", 100),
            )
        )
        return CharacterTextSplitter(
            chunk_size=options.splitter.chunk_size,
            chunk_overlap=options.splitter.chunk_overlap,
            separator="\n",
        )

    def split_content(self, content: str) -> List[str]:
        return content.split("\f")


class DocumentHandler(ABC, DocumentProcessingMixin):
    @abstractmethod
    def fetch_content(self, document: Document) -> str:
        pass

    def generate_metadata(self, document: Document) -> dict:
        # By default, consider the URL as the source.
        return {"source": document.url}

    def process(self, document: DocumentModel, dataset: Dataset) -> List[Document]:
        content = self.fetch_content(document)
        metadata = self.generate_metadata(document)
        pages = self.split_content(content)
        all_docs = []
        text_splitter = self.get_text_splitter(document)
        for page_content in pages:
            docs = [Document(page_content=page_content, metadata=metadata.copy())]
            all_docs.extend(docs)
        all_docs = text_splitter.split_documents(all_docs)
        document.page_size = len(all_docs)
        document.content_size = 0
        for segment in all_docs:
            document.content_size += len(segment.page_content)
        for page_number, doc in enumerate(all_docs):
            doc.metadata["page_number"] = page_number
            doc.metadata[
                "urn"
            ] = f"{dataset.id}-{document.url}-{doc.metadata['page_number']}"

        logger.info(
            f"got documents: {len(all_docs)} while loading dataset {dataset.id}"
        )
        return all_docs


class PDFHandler(DocumentHandler):
    @staticmethod
    def get_document_page_size(document: Document) -> int:
        if document.page_size != 0:
            return document.page_size
        else:
            pdf_content = GoogleCloudStorageClient().load(document.url)
            text = PDFHandler.extract_text_from_pdf(pdf_content)
            pages = text.split("\f")
            return len(pages)

    @staticmethod
    def extract_text_from_pdf(
        contents: io.BytesIO, preview_size: int = float("inf")
    ) -> list:
        pdf = pypdf.PdfReader(contents)
        num_pages = len(pdf.pages)
        total_text = []
        for page in range(min(num_pages, preview_size)):
            page_text = pdf.pages[page].extract_text()
            page_text = re.sub(r"\x01", " ", page_text)
            total_text.append(page_text)
        total_text = "\n".join(total_text)
        total_text = total_text.strip()
        return total_text

    def fetch_content(self, document: Document) -> str:
        storage_client = GoogleCloudStorageClient()
        pdf_content = storage_client.load(document.url)
        return self.extract_text_from_pdf(pdf_content)


class AnnotatedDataHandler(DocumentHandler):
    def fetch_content(self, document: Document) -> str:
        document.url = document.uid
        webhook_handler = AnnotatedDataStorageClient()
        return webhook_handler.load(document.uid)

    def split_content(self, content: str) -> List[str]:
        return [content]

    def generate_metadata(self, document: Document) -> dict:
        # For annotated data, the UID is the source.
        return {"source": document.uid}


class WordHandler(DocumentHandler):
    def fetch_content(
        self, document: Document, preview_size: int = float("inf")
    ) -> str:
        storage_client = GoogleCloudStorageClient()
        word_content = storage_client.load(document.url)
        preview_content_size = document.split_option["chunk_size"] * preview_size
        if document.url.endswith(".docx"):
            return self.extract_text_from_docx(word_content, preview_content_size)
        elif document.url.endswith(".doc"):
            return self.extract_text_from_doc(word_content, preview_content_size)
        else:
            raise Exception("Unsupported Word format")

    def extract_text_from_docx(
        self, content: io.BytesIO, preview_content_size: int = float("inf")
    ) -> str:
        word_doc = WordDocument(content)
        full_text = []
        for para in word_doc.paragraphs:
            full_text.append(para.text)
            if len(full_text) >= preview_content_size:
                break
        return "\n".join(full_text)

    def extract_text_from_doc(
        self, content: io.BytesIO, preview_content_size: int = float("inf")
    ) -> str:
        def get_unoconv_path() -> str:
            try:
                result = subprocess.run(
                    ["which", "unoconv"], capture_output=True, text=True
                )
                if result.returncode == 0:
                    return result.stdout.strip()
            except Exception as e:
                pass
            raise EnvironmentError("unoconv is not found in the system PATH.")

        unoconv_path = get_unoconv_path()
        # Convert .doc to .docx using unoconv
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_doc:
            temp_doc.write(content.read())
        temp_docx = temp_doc.name + "x"
        subprocess.run([unoconv_path, "-f", "docx", "-o", temp_docx, temp_doc.name])
        with open(temp_docx, "rb") as docx_file:
            full_text = self.extract_text_from_docx(docx_file, preview_content_size)
        os.remove(temp_doc.name)
        os.remove(temp_docx)
        return full_text


def load_and_split_documents(datasets: list[Dataset]):
    handlers = {
        "pdf": PDFHandler(),
        "annotated_data": AnnotatedDataHandler(),
        "word": WordHandler(),
    }

    docs = []
    # process each document
    for dataset in datasets:
        for document in dataset.documents:
            handler = handlers.get(document.type)
            if handler:
                processed_docs = handler.process(document, dataset)
                docs.extend(processed_docs)
            else:
                # Handle unsupported document types
                logger.error(f"Document type {document.type} not supported")
                raise Exception("Document type not supported")
    return docs
